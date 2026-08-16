#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Generate test files for batch 1 tool verification."""
import os

TEST_DIR = r'E:\网站项目\smartimgkit\_test_files'
os.makedirs(TEST_DIR, exist_ok=True)

# 1. Test PDF with text + table-like content
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import mm

pdf_path = os.path.join(TEST_DIR, 'test_sample.pdf')
c = canvas.Canvas(pdf_path, pagesize=A4)
w, h = A4
c.setFont('Helvetica-Bold', 16)
c.drawString(50, h - 50, 'Quarterly Sales Report')
c.setFont('Helvetica', 11)
c.drawString(50, h - 70, 'This is a test PDF for SmartImgKit conversion tools.')
# Table-like text (aligned columns)
c.setFont('Helvetica-Bold', 10)
c.drawString(50, h - 100, 'Product')
c.drawString(200, h - 100, 'Q1')
c.drawString(280, h - 100, 'Q2')
c.drawString(360, h - 100, 'Q3')
c.drawString(440, h - 100, 'Q4')
c.setFont('Helvetica', 10)
rows = [
    ('Widget A', '1200', '1500', '1800', '2100'),
    ('Widget B', '800', '950', '1100', '1300'),
    ('Widget C', '450', '520', '610', '700'),
    ('Gadget X', '2300', '2100', '2400', '2800'),
]
y = h - 120
for row in rows:
    c.drawString(50, y, row[0])
    c.drawString(200, y, row[1])
    c.drawString(280, y, row[2])
    c.drawString(360, y, row[3])
    c.drawString(440, y, row[4])
    y -= 18
c.drawString(50, y - 20, 'Total revenue increased by 23% year over year.')
# Second page
c.showPage()
c.setFont('Helvetica-Bold', 14)
c.drawString(50, h - 50, 'Page 2: Summary')
c.setFont('Helvetica', 11)
c.drawString(50, h - 70, 'The conversion tools should extract this text accurately.')
c.drawString(50, h - 90, 'Each paragraph becomes a separate line in the output.')
c.save()
print(f'OK PDF: {pdf_path} ({os.path.getsize(pdf_path)} bytes)')

# 2. Test DOCX with paragraphs and headings
from docx import Document
from docx.shared import Pt

docx_path = os.path.join(TEST_DIR, 'test_sample.docx')
doc = Document()
doc.add_heading('Test Document for Word to PDF', level=1)
doc.add_paragraph('This is the first paragraph. It contains regular text that should be converted to PDF.')
doc.add_heading('Section One', level=2)
doc.add_paragraph('The quick brown fox jumps over the lazy dog. This sentence is used to test text rendering.')
doc.add_paragraph('Another paragraph with different content. The PDF should preserve paragraph breaks.')
doc.add_heading('Section Two', level=2)
doc.add_paragraph('Bullet points and lists may not be fully preserved, but text content will be accurate.')
doc.add_paragraph('Final paragraph. The conversion is 100% browser-based with no server uploads.')
doc.save(docx_path)
print(f'OK DOCX: {docx_path} ({os.path.getsize(docx_path)} bytes)')

# 3. Test XLSX with a table
import openpyxl

xlsx_path = os.path.join(TEST_DIR, 'test_sample.xlsx')
wb = openpyxl.Workbook()
ws = wb.active
ws.title = 'Sales'
ws.append(['Product', 'Q1', 'Q2', 'Q3', 'Q4', 'Total'])
data = [
    ['Widget A', 1200, 1500, 1800, 2100, 6600],
    ['Widget B', 800, 950, 1100, 1300, 4150],
    ['Widget C', 450, 520, 610, 700, 2280],
    ['Gadget X', 2300, 2100, 2400, 2800, 9600],
    ['Gadget Y', 1100, 1300, 1500, 1700, 5600],
]
for row in data:
    ws.append(row)
# Second sheet
ws2 = wb.create_sheet('Summary')
ws2.append(['Metric', 'Value'])
ws2.append(['Total Products', 5])
ws2.append(['Total Revenue', 28230])
ws2.append(['Best Seller', 'Gadget X'])
wb.save(xlsx_path)
print(f'OK XLSX: {xlsx_path} ({os.path.getsize(xlsx_path)} bytes)')

print('\nAll test files created in', TEST_DIR)
